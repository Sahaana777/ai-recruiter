"""
resume_parser.py
=================
Utility to turn a resume file (PDF / DOCX / TXT) into plain text so it can
be fed into the same extraction pipeline used for conversational input.

This keeps Part 1 (extraction) and Part 2 (resume parsing + matching)
sharing one source of truth for what counts as a "skill", "technology",
or "language" -- we don't want two different vocabularies drifting apart.
"""

from pathlib import Path

import pdfplumber
import docx


class UnsupportedFileType(ValueError):
    pass


def parse_resume(file_path: str) -> str:
    """Extract raw text from a resume file. Supports .pdf, .docx, .txt"""
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)
    elif ext == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Use .pdf, .docx, or .txt"
        )


def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:

    import io

    ext = Path(filename).suffix.lower()
    buffer = io.BytesIO(file_bytes)

    if ext == ".pdf":
        text_parts = []
        with pdfplumber.open(buffer) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text_parts.append(t)
        return "\n".join(text_parts)
    elif ext == ".docx":
        document = docx.Document(buffer)
        return "\n".join(p.text for p in document.paragraphs)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise UnsupportedFileType(
            f"Unsupported file type '{ext}'. Use .pdf, .docx, or .txt"
        )


def _parse_pdf(path: Path) -> str:
    text_parts = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return "\n".join(text_parts)


def _parse_docx(path: Path) -> str:
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)
